#!/usr/bin/env python3
"""Build the Sayari FDE submission report (.docx).

Reflects the current state of the project, including:
  - list_1 ownership-gap story (silent miss)
  - list_3 false positive story (noisy collision)
  - The deterministic-engine + grounded-AI architecture

Run from the repo root:
    python scripts/build_report.py

Requires python-docx:
    pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_REPO = Path(__file__).resolve().parents[1]
DEST = _REPO / "Meridian_Sentinel_FDE_Report.docx"

NAVY = "1F3A5F"
GRAY = "7F8C8D"
ACCENT = "8A6D1F"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.8)
sec.left_margin = sec.right_margin = Inches(0.9)
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)


def setc(run, hexc):
    run.font.color.rgb = RGBColor.from_string(hexc)


def heading(text, size=13, color=NAVY, before=10, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    setc(r, color)
    return p


def body(text, after=6, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def bullet(lead, rest, after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if lead:
        r = p.add_run(lead)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(rest)
    r2.font.size = Pt(10.5)
    return p


# ── Title block ──────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(1)
r = t.add_run("Meridian Sentinel: A Trustworthy AI Counterparty Screening Copilot Built on Sayari")
r.bold = True
r.font.size = Pt(16)
setc(r, NAVY)

s = doc.add_paragraph()
s.paragraph_format.space_after = Pt(2)
r = s.add_run("Forward Deployed Engineer Technical Exercise · Scenario 2 (Analytics), extended with Scenario 1 (external source enrichment)")
r.font.size = Pt(10)
setc(r, GRAY)

s2 = doc.add_paragraph()
s2.paragraph_format.space_after = Pt(8)
r = s2.add_run("Peter Nemrow · June 2026 · Datasets: list_1 (50 high risk / state owned entities) and list_3 (50 supply chain auto parts vendors)")
r.font.size = Pt(9.5)
setc(r, GRAY)

pPr = s2._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "2E75B6")
pbdr.append(bottom)
pPr.append(pbdr)

# ── 1. The engagement ────────────────────────────────────────────────────────
heading("The engagement")
body(
    "This is a proof of concept for Meridian Energy Trading SA, a Geneva commodities trader that must vet vendors "
    "and counterparties before onboarding them. Their status quo, a name based screen against the OFAC SDN list, has "
    "two structural blind spots. First, it cannot flag an unlisted company that is ≥50% owned or controlled by a "
    "sanctioned party, even though OFAC's 50% rule (31 CFR §501.801) prohibits transacting with all the same. Second, "
    "it fires noisily on legitimate vendors whose names collide with sanctioned entities, consuming analyst review "
    "time and eroding trust in the screen. Closing both gaps is the point of the tool."
)

# ── 2. Approach ──────────────────────────────────────────────────────────────
heading("Approach")
body(
    "I chose Scenario 2 (analytics) and extended it with Scenario 1 (external source enrichment) by screening every "
    "entity against the live OFAC SDN feed. A deterministic engine sits at the core: resolve each name to a Sayari "
    "entity, fetch its profile, traverse its ownership graph, aggregate. No LLM in the data path, so every figure is "
    "reproducible and traceable to a specific Sayari field. That engine is exposed as typed tools (a FastAPI service "
    "and an MCP server), and a grounded, streaming copilot sits on top: it may only state facts a tool returned, "
    "cites every claim, and flags anything uncertain. A web console is the analyst's delivery vehicle. The guiding "
    "principle is Sayari's own: AI that shows its work and traces every finding to its source."
)

# ── 3. The headline finding ──────────────────────────────────────────────────
heading("The headline finding")
body(
    "Of the 50 entities in list_1, 49 resolved (98%) and 45 are sanctioned. The decisive result is the comparison "
    "against a fair OFAC name screen: it catches 33 of 40 OFAC exposed entities and misses 7. Four are hidden behind "
    "ownership (the 50% rule) and three are lost to name variations (transliteration or legal name differences). "
    "Sayari catches all 40. The clearest case is Russian Railways: absent from the SDN by name, yet controlled by "
    "sanctioned parties through the Russian state. Exposure no name screen can see.",
    after=6,
)
body(
    "The complementary story is in list_3, a 50 row sample of supply chain auto parts vendors. Here Sayari's "
    "resolution disambiguates four legitimate companies from sanctioned namesakes the OFAC name screen flags. Magna "
    "International matches a Mexican drug trafficker (SDN 6866). Continental matches an unrelated SDN entity (SDN "
    "54200). NSK matches an initialism collision (SDN 47854). Mando matches a drug trafficker's alias (SDN 54225). "
    "Sayari resolves all four correctly to the legitimate corporate entities. The lesson is symmetric: name screening "
    "fails in two distinct ways, silent miss and noisy false positive, and resolution plus the ownership graph closes "
    "both.",
    after=8,
)

# ── Reconciliation table (list_1) ────────────────────────────────────────────
rows = [
    ("Outcome", "n", "Meaning"),
    ("Both caught", "33", "On the SDN by name; a fair screen finds it (agreement)."),
    ("Sayari only (ownership)", "2", "Not named; flagged via ≥50% ownership (the gap)."),
    ("Screen hit wrong party", "2", "Screen fired on a different SDN entity; real entity exposed via ownership (the gap)."),
    ("Screen missed (name)", "3", "On the SDN, but the name did not match; resolution catches it."),
    ("OFAC only (review)", "2", "Screen hit; surfaced for review (incl. a verified Sayari linkage gap)."),
    ("No OFAC exposure", "7", "No OFAC SDN finding (some EU / UK sanctioned)."),
]
tbl = doc.add_table(rows=len(rows), cols=3)
tbl.style = "Table Grid"
tbl.columns[0].width = Inches(2.1)
tbl.columns[1].width = Inches(0.4)
tbl.columns[2].width = Inches(4.2)
for ri, (a, b, c) in enumerate(rows):
    cells = tbl.rows[ri].cells
    for ci, val in enumerate((a, b, c)):
        cells[ci].width = [Inches(2.1), Inches(0.4), Inches(4.2)][ci]
        para = cells[ci].paragraphs[0]
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(val)
        run.font.size = Pt(9)
        if ri == 0:
            run.bold = True
            setc(run, "FFFFFF")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), NAVY)
            cells[ci]._tc.get_or_add_tcPr().append(shd)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 4. Assumptions ───────────────────────────────────────────────────────────
heading("Assumptions")
bullet("Datasets: ", "list_1 (50 high risk / state owned entities) is the primary set for the silent miss story; list_3 (50 supply chain auto parts vendors) is the secondary set for the false positive story. The pipeline is list agnostic and runs unchanged on any uploaded vendor list.")
bullet("Resolution: ", "Sayari's best match is accepted, then marquee entities are verified by relationship degree and registration / tax identifiers. Four parent IDs are pinned with audit comments after manual verification.")
bullet("Geography: ", "'Headquarters' comes from the input list; 'exposure' comes from the ownership graph. Two complementary views.")
bullet("Point in time: ", "Sanctions status reflects the API at run time (May 2026). The OFAC baseline is a fair, transliteration aware screen. Demo 'cached' mode replays recorded real runs verbatim.")

# ── 5. Challenges ────────────────────────────────────────────────────────────
heading("Challenges and how I addressed them")
bullet(
    "Entity resolution is the make or break step. ",
    "Matching on the supplied street addresses surfaced subsidiaries over parents. 'Sberbank' resolved to a back office LLC, 'VTB Bank' to VTB Capital Holdings. I caught these with a resolution audit log and a transliteration aware mismatch flag, verified the true parent by relationship degree, and pinned three confirmed IDs. The logic generalizes into a parent preference heuristic with surfaced confidence flags so it holds on unseen lists.",
)
bullet(
    "Keeping the comparison honest. ",
    "An early version inadvertently flattered Sayari by under powering the OFAC screen. I rebuilt a fair, transliteration aware name screen, which made it catch more entities and made Sayari's advantage rest on the structural 50% rule gap rather than a handicapped baseline.",
)
bullet(
    "Surfacing disagreements, including a real one. ",
    "By cross referencing registration and tax IDs against the SDN, the tool flagged Kalashnikov Concern as conclusively the SDN listed entity yet missing the entity level OFAC factor on Sayari's parent record (it is present on its subsidiaries). The tool surfaces this for analyst review rather than hiding it: reconciliation, not a verdict.",
)
bullet(
    "Trust by construction. ",
    "No model touches the numbers; every value cites its raw field; field shapes were confirmed against live API responses rather than documentation. The grounded copilot calls the same typed tools the rest of the UI uses; the system prompt forbids fabrication.",
)

# ── 6. Value ─────────────────────────────────────────────────────────────────
heading("How this demonstrates the value of Sayari data")
body(
    "Meridian Sentinel is a reconciliation layer across three independent signals (an OFAC name screen, Sayari's risk "
    "factors, and Sayari's ownership graph) that surfaces where they disagree for a human to adjudicate, with every "
    "finding cited and therefore defensible to a regulator. It addresses both failure modes of name based screening "
    "demonstrably: the silent miss (list_1) and the noisy false positive (list_3). And it fits the real workflow: "
    "counterparties arrive via SFTP or a data warehouse, analysts disposition each one (maker checker), and a "
    "persisted, machine readable result payload flows back into the client's systems via API."
)

# ── 7. Real vs representative ────────────────────────────────────────────────
heading("What's real, what's representative, and what's next", size=12, color=GRAY)
body(
    "Real and verifiable: the deterministic engine, resolution with verified pins, the OFAC reconciliation across "
    "both datasets, the ownership graph, the grounded copilot, file based persistence with per run scoping, briefing "
    "PDFs, and the persisted result payload / API. Representative (designed, not wired): authentication, the "
    "integrations catalog beyond SFTP, and the Postgres path (the schema exists and the code path falls back to file "
    "storage gracefully). With more time: fully generalize parent preference resolution, deepen ownership traversal, "
    "add a hallucination rate evaluation harness, and ship a customer specific deployment harness.",
    after=4,
)

f = doc.add_paragraph()
f.paragraph_format.space_before = Pt(6)
r = f.add_run("Repository: <link> · Demo recording: <link> · Deterministic engine: packages/engine/")
r.font.size = Pt(9)
setc(r, GRAY)

doc.save(str(DEST))
print(f"saved {DEST}")
print(f"paragraphs: {len(doc.paragraphs)} | tables: {len(doc.tables)}")
